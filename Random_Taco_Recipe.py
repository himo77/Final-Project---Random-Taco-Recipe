# import modules
from PIL import Image, ImageDraw, ImageFont
import requests
import docx

image = Image.open('tai-s-captures-JiRSy0GfqPA-unsplash.jpg')   # open an image of type Image

width = image.width    # save the width of the image
height = image.height  # save the height of the image

print(width, height)   # print the width and height of the image

image.thumbnail((800, 800))  # transform the image to a thumbnail

img_draw = ImageDraw.Draw(image)  # create an object to use to draw on the image

font = ImageFont.truetype('DejaVuSans.ttf', 30)  # create a font
img_draw.text([120, 460], 'Random Taco Cookbook', fill='purple', font=font)  # write on the image

image.show()  # display the image
image.save('Random Taco.jpg')  # save the image

taco_list = []  # create an empty list

for item in range(3):  # loop 3 times
    try:  # treat errors
        # make get request to collect data in json format
        taco_url = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()
        taco_list.append(taco_url)  # add the collected json data to the end of the list
        print(taco_list)  # print the list of data
    except:  # catch exception to deal with it
        print('stop')

document = docx.Document()  # create an object document to build word document


document.add_paragraph('Random Taco Cookbook', 'Title')   # Add a title

document.add_picture('Random Taco.jpg', width=docx.shared.Inches(6), height=docx.shared.Cm(12))  # Add a picture

document.add_paragraph('Credits', 'Heading 1')   # Add a Heading


document.add_paragraph('Taco image: Photo by Tai\'s Captures on Unsplash')  # Add a Paragraph to who taken the image

document.add_paragraph('Tacos from: https://taco-1150.herokuapp.com/random/?full_taco=true')  # Add a Paragraph of the source of the recipe

document.add_paragraph('Code by: Ibrahim Hamadon')   # Add a paragraph with the coder name

document.add_page_break()  # Add a break page


for item in range(len(taco_list)):   # loop on the list of collected data to create list of names and recipes
    document.add_paragraph(taco_list[item]["base_layer"]["name"], 'heading 1')  # add a heading
    document.add_paragraph(taco_list[item]["base_layer"]["recipe"])
    document.add_paragraph(taco_list[item]["seasoning"]["name"], 'heading 1')  # add a heading
    document.add_paragraph(taco_list[item]["seasoning"]["recipe"])
    document.add_paragraph(taco_list[item]["mixin"]["name"], 'heading 1')  # add a heading
    document.add_paragraph(taco_list[item]["mixin"]["recipe"])
    document.add_paragraph(taco_list[item]["condiment"]["name"], 'heading 1')  # add a heading
    document.add_paragraph(taco_list[item]["condiment"]["recipe"])
    document.add_paragraph(taco_list[item]["shell"]["name"], 'heading 1')  # add a heading
    document.add_paragraph(taco_list[item]["shell"]["recipe"])

    document.add_page_break()  # add break page

document.save('taco_recipe.docx')  # save the document











